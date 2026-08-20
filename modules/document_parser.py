from __future__ import annotations

from pathlib import Path
from typing import Any


def extract_uploaded_text(uploaded_file: Any) -> str:
    """Extract text from Streamlit uploads without persisting the source file."""
    if uploaded_file is None:
        return ""
    name = Path(getattr(uploaded_file, "name", "upload.txt")).name
    suffix = Path(name).suffix.lower()
    raw = uploaded_file.getvalue()
    if suffix in {".txt", ".md", ".csv"}:
        return raw.decode("utf-8", errors="ignore").strip()
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(raw)
            return "\n".join(page.extract_text() or "" for page in reader.pages).strip()
        except Exception as exc:
            raise ValueError(f"Could not read the PDF file: {exc}") from exc
    if suffix == ".docx":
        try:
            from docx import Document
            from io import BytesIO

            document = Document(BytesIO(raw))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                paragraphs.extend(" | ".join(cell.text for cell in row.cells) for row in table.rows)
            return "\n".join(line for line in paragraphs if line.strip()).strip()
        except Exception as exc:
            raise ValueError(f"Could not read the DOCX file: {exc}") from exc
    raise ValueError("Unsupported file type. Upload PDF, DOCX, TXT, MD, or CSV.")


def compact_text(text: str, limit: int = 12000) -> str:
    """Normalize whitespace and cap prompt size for predictable model calls."""
    normalized = " ".join((text or "").split())
    return normalized[:limit]


def file_label(uploaded_file: Any) -> str:
    return Path(getattr(uploaded_file, "name", "uploaded document")).name if uploaded_file else "Not uploaded"
